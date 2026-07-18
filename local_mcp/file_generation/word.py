"""Markdown-like content -> Word (.docx) renderer."""

from __future__ import annotations

import re
from pathlib import Path

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
_RULE_CHARS = {"-", "_", "*"}


def render_docx_file(content: str, target: Path, *, title: str = "") -> int:
    """Render Markdown-like text into a .docx file and return the bytes written."""
    document = _new_document()
    source = content or ""
    if not source.strip() and title:
        source = f"# {title}"
    _render_blocks(document, source)
    document.save(str(target))
    return target.stat().st_size


def _new_document():
    try:
        from docx import Document
    except ImportError as err:
        raise ValueError(
            "Word output requires the python-docx package. Install it with: pip install python-docx"
        ) from err
    return Document()


def _render_blocks(document, content: str) -> None:
    paragraph_buffer: list[str] = []
    code_lines: list[str] = []
    in_code_block = False

    def flush_paragraph() -> None:
        if paragraph_buffer:
            _add_inline_runs(document.add_paragraph(), " ".join(paragraph_buffer))
            paragraph_buffer.clear()

    for raw_line in content.splitlines():
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                _add_code_paragraph(document, code_lines)
                code_lines = []
            in_code_block = not in_code_block
            continue
        if in_code_block:
            code_lines.append(raw_line.rstrip())
            continue
        if not stripped:
            flush_paragraph()
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            flush_paragraph()
            document.add_heading(_plain_text(heading.group(2)), level=min(len(heading.group(1)), 4))
            continue
        if len(stripped) >= 3 and set(stripped) <= _RULE_CHARS:
            flush_paragraph()
            continue
        bullet = _BULLET_RE.match(raw_line)
        if bullet:
            flush_paragraph()
            _add_inline_runs(document.add_paragraph(style="List Bullet"), bullet.group(1))
            continue
        numbered = _NUMBERED_RE.match(raw_line)
        if numbered:
            flush_paragraph()
            _add_inline_runs(document.add_paragraph(style="List Number"), numbered.group(1))
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            _add_inline_runs(document.add_paragraph(style="Intense Quote"), stripped.lstrip("> "))
            continue
        paragraph_buffer.append(stripped)

    if code_lines:
        _add_code_paragraph(document, code_lines)
    flush_paragraph()


def _add_code_paragraph(document, code_lines: list[str]) -> None:
    from docx.shared import Pt

    run = document.add_paragraph().add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_inline_runs(paragraph, text: str) -> None:
    for token in _INLINE_TOKEN_RE.split(_LINK_RE.sub(r"\1 (\2)", text)):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            paragraph.add_run(token[1:-1]).font.name = "Courier New"
        else:
            paragraph.add_run(token)


def _plain_text(text: str) -> str:
    return _LINK_RE.sub(r"\1", text).replace("**", "").replace("`", "").strip()
